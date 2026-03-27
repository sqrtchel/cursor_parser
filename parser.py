import sys
import argparse
import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import docx
from tempfile import NamedTemporaryFile
import re
import pyodbc
from dotenv import load_dotenv
import logging
from logging.handlers import RotatingFileHandler
import pandas as pd
import torch
from sentence_transformers import SentenceTransformer, util
import numpy as np

#Настройка логирования
log_dir = "ParserLogs"
if not os.path.exists(log_dir):
    os.makedirs(log_dir)

log_file = os.path.join(log_dir, "parser.log")
handler = RotatingFileHandler(log_file, maxBytes=1*1024*1024, backupCount=3, encoding='utf-8')

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    handlers=[handler, logging.StreamHandler()] # Добавляем вывод в консоль
)


# Загружаем переменные окружения
load_dotenv()


class SemanticAnalyzer:
    """Класс для интеллектуального анализа текста на основе семантической схожести."""
    
    def __init__(self, excel_path, cache_path="embeddings_cache.pt"):
        logging.info("Инициализация семантического анализатора...")
        # Используем модель, которая понимает русский язык
        self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        self.reference_embeddings = {}
        self.excel_path = excel_path
        self.cache_path = cache_path
        self._load_data()

    def _load_data(self):
        """Загрузка данных: сначала из кеша, если нет - из Excel."""
        if os.path.exists(self.cache_path):
            try:
                logging.info(f"Загрузка векторов из кеша: {self.cache_path}...")
                self.reference_embeddings = torch.load(self.cache_path, weights_only=False)
                logging.info("Векторы успешно загружены из кеша.")
                return
            except Exception as e:
                logging.error(f"Ошибка при загрузке кеша: {e}. Попробуем загрузить из Excel.")

        # Если кеша нет или ошибка загрузки - обучаемся на Excel
        self._load_from_excel()

    def _load_from_excel(self):
        """Загрузка данных из Excel и создание эталонных векторов."""
        try:
            if not os.path.exists(self.excel_path):
                logging.error(f"Файл с данными {self.excel_path} не найден. Семантический поиск будет работать в упрощенном режиме.")
                return

            logging.info(f"Читаем эталонную базу из {self.excel_path}...")
            df = pd.read_excel(self.excel_path)
            logging.info(f"Загружено {len(df)} строк из эталонной базы.")

            # Карта соответствия колонок в Excel и групп в парсере
            column_map = {
                "ИСТОЧНИК ФИНАНСИРОВАНИЯ": "Источник финансирования",
                "ТРЕБОВАНИЯ К СРОКУ ГОДНОСТИ": "Требования к сроку годности",
                "ПОРЯДОК ОПЛАТЫ ТОВАРОВ": "Порядок оплаты товаров",
                "УСЛОВИЯ ПОСТАВКИ": "Условия поставки",
                "Примечание к сроку действия контракта": "Примечание к сроку действия контракта",
                "СРОК ДЕЙСТВИЯ КОНТРАКТА": "Срок действия контракта"
            }

            for col, group in column_map.items():
                if col in df.columns:
                    # Берем уникальные непустые значения для каждой группы
                    examples = df[col].dropna().unique().tolist()
                    if examples:
                        logging.info(f"Создание векторов для группы '{group}' ({len(examples)} примеров)...")
                        # Ограничиваем количество примеров для скорости (берем 2000 самых разнообразных или просто первых)
                        examples = examples[:2000] 
                        embeddings = self.model.encode(examples, convert_to_tensor=True)
                        self.reference_embeddings[group] = embeddings
                else:
                    logging.warning(f"Колонка {col} не найдена в Excel.")

            # Сохраняем в кеш после успешного создания
            if self.reference_embeddings:
                self._save_to_cache()

        except Exception as e:
            logging.error(f"Ошибка при обучении семантического анализатора: {e}")

    def _save_to_cache(self):
        """Сохранение векторов в файл кеша."""
        try:
            logging.info(f"Сохранение векторов в кеш: {self.cache_path}...")
            torch.save(self.reference_embeddings, self.cache_path)
            logging.info("Кеш успешно создан.")
        except Exception as e:
            logging.error(f"Не удалось сохранить кеш: {e}")

    def find_best_sentence(self, sentences, group_name, threshold=0.91):
        """Находит наиболее подходящее предложение для указанной группы."""
        if not sentences or group_name not in self.reference_embeddings:
            return None

        # Кодируем все предложения документа
        sentence_embeddings = self.model.encode(sentences, convert_to_tensor=True)
        
        # Считаем схожесть каждого предложения со всеми эталонами группы
        cos_scores = util.cos_sim(sentence_embeddings, self.reference_embeddings[group_name])
        
        # Для каждого предложения берем максимальную схожесть с любым из эталонов
        max_scores, _ = torch.max(cos_scores, dim=1)
        
        # Находим лучший результат
        best_idx = torch.argmax(max_scores).item()
        best_score = max_scores[best_idx].item()

        if best_score >= threshold:
            logging.info(f"Найдено совпадение для '{group_name}' (Score: {best_score:.2f}): {sentences[best_idx][:100]}...")
            return sentences[best_idx]
        
        logging.info(f"Для группы '{group_name}' ни одно предложение не прошло порог {threshold}.")
        return None


class DocumentParser:
    def __init__(self, semantic_analyzer=None):
        self.groups_config = {
            "Источник финансирования": ["источник финансирования", "бюджетные средства", "внебюджетные средства", "средства бюджета", "собственных средств", "финансирование"],
            "Требования к сроку годности": ["срок годности", "остаточный срок годности", "дата изготовления", "гарантийный срок"],
            "Порядок оплаты товаров": ["порядок оплаты", "условия оплаты", "безналичный расчет", "аванс", "платеж", "банковские реквизиты", "в течение"],
            "Условия поставки": ["условия поставки", "место поставки", "срок поставки", "доставка", "транспортировка", "пункт назначения", "партиями", "заявке"],
            "Примечание к сроку действия контракта": ["контракт вступает в силу", "действие контракта", "действует до", "срок действия контракта"],
            "Срок действия контракта": ["срок действия", "до полного исполнения", "дата окончания действия"]
        }

        self.groups = {group: [] for group in self.groups_config.keys()}
        self.contract_date = None
        self.semantic_analyzer = semantic_analyzer


    def parse_by_number(self, number, provided_url=None):
        # метод принимает номер закупки и необязательную ссылку на общую информацию


        logging.info(f"Начинаем обработку закупки: {number}")
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

        url_documents = None
        
        # Определяем ФЗ по номеру
        is_223fz = len(str(number)) == 11 and str(number).startswith('3')
        
        # Пытаемся извлечь ID из предоставленной ссылки или через поиск
        if provided_url:
            notice_match = re.search(r'noticeInfoId=(\d+)', provided_url)
            reg_match = re.search(r'regNumber=(\d+)', provided_url)
            
            # Пытаемся вытащить тип извещения (ea20, ep44 и тд)
            type_match = re.search(r'notice/([^/]+)/', provided_url)
            notice_type = type_match.group(1) if type_match else "ea20"
            
            if is_223fz:
                if notice_match:
                    url_documents = f"https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?noticeInfoId={notice_match.group(1)}"
                else:
                    notice_id = self._get_223fz_notice_id(number)
                    if notice_id:
                        url_documents = f"https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?noticeInfoId={notice_id}"
            else: # 44-ФЗ
                if reg_match:
                    url_documents = f"https://zakupki.gov.ru/epz/order/notice/{notice_type}/view/documents.html?regNumber={reg_match.group(1)}"
        
        # Если ссылка на документы всё еще не сформирована, строим дефолтную по номеру
        if not url_documents:
            if is_223fz:
                notice_id = self._get_223fz_notice_id(number)
                if notice_id:
                    url_documents = f"https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?noticeInfoId={notice_id}"
            else:
                # Для 44-ФЗ по умолчанию ea20, но это может не сработать для всех
                url_documents = f"https://zakupki.gov.ru/epz/order/notice/ea20/view/documents.html?regNumber={number}"

        # Парсинг документов
        if url_documents:
            try:
                logging.info(f"Загружаем документы по ссылке: {url_documents}")
                response_docs = requests.get(url_documents, headers=headers, timeout=15)
                response_docs.raise_for_status()
                soup_docs = BeautifulSoup(response_docs.text, 'html.parser')

                # ищем ссылки на все поддерживаемые форматы
                doc_links = self._find_supported_links(soup_docs, url_documents)
                if doc_links:
                    logging.info(f"Найдено {len(doc_links)} файлов для анализа.")
                    all_paragraphs = []
                    for link in doc_links:
                        paragraphs = self._read_file_from_url(link)
                        all_paragraphs.extend(paragraphs)
                    
                    # Группируем текст из всех найденных документов
                    self._group_paragraphs(all_paragraphs)
                else:
                    logging.warning(f"Файлы для анализа не найдены по ссылке {url_documents}.")

            except Exception as e:
                logging.error(f"Ошибка при загрузке документов: {e}")
        else:
            logging.warning(f"Ссылка на документы не была сформирована для закупки {number}.")

        self._process_contract_date()
        return self.groups

    def _find_supported_links(self, soup, base_url):
        # Метод ищет ссылки на docx, doc, pdf, xlsx
        links = set()
        extensions = ['.docx', '.doc', '.pdf', '.xlsx']
        ignore_patterns = [
            "электронный документ, полученный из внешней системы",
            "zakupki-traffic",
            "traffic",
            "печатная форма"
        ]

        for link in soup.find_all('a'):
            href = link.get('href', '')
            title = link.get('title', '')
            text = link.get_text().strip()

            # Игнорируем файлы по названию или ссылке
            full_text = f"{text.lower()} {title.lower()} {href.lower()}"
            if any(pattern in full_text for pattern in ignore_patterns):
                continue

            is_supported = any(ext in href.lower() or ext in title.lower() or ext in text.lower() for ext in extensions)
            
            if (is_supported or 'downloaddoc' in href.lower() or 'filestore' in href.lower()) and href:
                full_url = urljoin(base_url, href)
                links.add(full_url)

        logging.info(f"Итого найдено потенциальных файлов для анализа: {len(links)}")
        return list(links)

    def _read_file_from_url(self, url):
        # Универсальный метод скачивания и чтения файлов разных форматов
        paragraphs = []
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                logging.info(f"Загрузка файла (попытка {attempt + 1}): {url}")
                headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                # Увеличиваем таймаут для стабильности
                response = requests.get(url, headers=headers, stream=True, timeout=20)
                response.raise_for_status()

                # Проверяем заголовки для определения реального имени файла
                content_disposition = response.headers.get('Content-Disposition', '')
                filename_from_header = ""
                if 'filename=' in content_disposition:
                    # Извлекаем имя файла из заголовка (может быть в кавычках)
                    filename_from_header = re.findall('filename="?([^";]+)"?', content_disposition)
                    filename_from_header = filename_from_header[0] if filename_from_header else ""

                # Проверяем сигнатуры файлов в начале контента
                file_content = response.content
                if not file_content:
                    logging.warning(f"Файл {url} пуст.")
                    return []

                # Проверяем, не HTML ли это (иногда ссылки ведут на страницу ошибки/логина)
                if file_content.strip().startswith(b'<!DOCTYPE') or file_content.strip().startswith(b'<html'):
                    logging.warning(f"Файл {url} является HTML страницей, а не документом. Пропускаем.")
                    return []

                with NamedTemporaryFile(delete=False) as tmp_file:
                    tmp_file.write(file_content)
                    tmp_path = tmp_file.name

                # Определяем расширение
                ext_source = (filename_from_header or url).lower()
                
                # Попытка чтения на основе сигнатуры и расширения
                if file_content.startswith(b'PK\x03\x04'): # zip/docx/xlsx
                    # Сначала проверяем, не является ли это обычным zip-архивом (часто в 44-ФЗ)
                    paragraphs = self._read_zip_archive(tmp_path)
                    
                    if not paragraphs:
                        if '.xlsx' in ext_source:
                            paragraphs = self._read_xlsx(tmp_path, silent=True)
                        elif '.docx' in ext_source:
                            paragraphs = self._read_docx(tmp_path, silent=True)
                        else:
                            # Если расширения нет, пробуем оба варианта
                            paragraphs = self._read_docx(tmp_path, silent=True)
                            if not paragraphs:
                                paragraphs = self._read_xlsx(tmp_path, silent=True)
                            
                elif file_content.startswith(b'%PDF'):
                    paragraphs = self._read_pdf(tmp_path)
                elif file_content.startswith(b'\xd0\xcf\x11\xe0'): # Старый DOC/XLS (OLE2)
                    paragraphs = self._read_doc(tmp_path)
                else:
                    # Если сигнатура неизвестна, пробуем по расширению
                    if '.docx' in ext_source:
                        paragraphs = self._read_docx(tmp_path)
                    elif '.doc' in ext_source:
                        paragraphs = self._read_doc(tmp_path)
                    elif '.pdf' in ext_source:
                        paragraphs = self._read_pdf(tmp_path)
                    elif '.xlsx' in ext_source:
                        paragraphs = self._read_xlsx(tmp_path)

                os.unlink(tmp_path)
                if paragraphs:
                    logging.info(f"Успешно извлечено {len(paragraphs)} строк из {url}")
                    return paragraphs
                else:
                    # Если расширение было .doc, выводим более понятное сообщение
                    if '.doc' in ext_source and not '.docx' in ext_source:
                        logging.warning(f"Файл {url} (старый .doc) не может быть прочитан без установленного MS Word. Пропускаем.")
                    else:
                        logging.warning(f"Не удалось извлечь текст из {url} (формат не распознан или файл пуст).")
                    return []

            except requests.exceptions.RequestException as e:
                logging.error(f"Ошибка сети при загрузке {url} (попытка {attempt + 1}): {e}")
                if attempt == max_retries - 1:
                    return []
            except Exception as e:
                logging.error(f"Критическая ошибка при обработке файла {url}: {e}")
                return []

        return paragraphs

    def _read_docx(self, path, silent=False):
        paragraphs = []
        try:
            doc = docx.Document(path)
            # Извлекаем текст из обычных абзацев
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Дополнительно извлекаем текст из всех таблиц в документе
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # В одной ячейке может быть несколько абзацев текста
                        cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        # Объединяем ячейки строки в одну строку для анализа
                        paragraphs.append(" | ".join(row_data))
                        
        except Exception as e:
            if not silent:
                logging.error(f"Ошибка чтения DOCX: {e}")
        return paragraphs

    def _read_doc(self, path):
        # Попытка прочитать .doc файл. На Windows это требует MS Word.
        paragraphs = []
        try:
            # На Windows можно попробовать через win32com, если Word установлен
            import win32com.client as win32
            import pythoncom
            
            # Инициализация COM-библиотеки для работы в многопоточном режиме
            pythoncom.CoInitialize()
            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            
            # Получаем абсолютный путь
            abs_path = os.path.abspath(path)
            doc = word.Documents.Open(abs_path)
            
            # Извлекаем текст
            full_text = doc.Content.Text
            paragraphs = [p.strip() for p in full_text.split('\r') if p.strip()]
            
            doc.Close(False)
            word.Quit()
        except Exception as e:
            # Если Word не установлен или произошла ошибка, paragraphs останется пустым
            pass
            
        return paragraphs

    def _read_pdf(self, path):
        paragraphs = []
        try:
            import fitz # PyMuPDF
            doc = fitz.open(path)
            for page in doc:
                # 1. Извлекаем обычный текст
                text = page.get_text()
                # Очистка текста от лишних переносов внутри слов и строк
                text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                paragraphs.extend(lines)
                
                # 2. Пытаемся найти таблицы (если есть)
                try:
                    tabs = page.find_tables()
                    for table in tabs:
                        for row in table.extract():
                            # Очищаем ячейки от None и лишних пробелов
                            row_data = [str(cell).strip() for cell in row if cell is not None]
                            if row_data:
                                paragraphs.append(" | ".join(row_data))
                except Exception:
                    # Некоторые PDF могут не поддерживать поиск таблиц или быть защищены
                    pass
                    
            doc.close()
        except ImportError:
            logging.error("Библиотека PyMuPDF (fitz) не установлена.")
        except Exception as e:
            logging.error(f"Ошибка чтения PDF: {e}")
        return paragraphs

    def _read_xlsx(self, path, silent=False):
        paragraphs = []
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell).strip() for cell in row if cell is not None])
                    if row_text:
                        paragraphs.append(row_text)
        except ImportError:
            if not silent:
                logging.error("Библиотека openpyxl не установлена.")
        except Exception as e:
            if not silent:
                logging.error(f"Ошибка чтения XLSX: {e}")
        return paragraphs

    def _read_zip_archive(self, path):
        # Метод проверяет, не является ли файл обычным zip-архивом с документами внутри
        import zipfile
        paragraphs = []
        
        try:
            with zipfile.ZipFile(path, 'r') as zf:
                # Список файлов в архиве
                file_list = zf.namelist()
                
                # Если в архиве есть папка 'word/' или '[Content_Types].xml', 
                # то это скорее всего docx/xlsx файл, а не архив документов.
                if 'word/document.xml' in file_list or 'xl/workbook.xml' in file_list:
                    return []
                
                logging.info(f"Обнаружен ZIP-архив, содержащий {len(file_list)} файлов.")
                for filename in file_list:
                    if any(filename.lower().endswith(ext) for ext in ['.docx', '.doc', '.pdf', '.xlsx']):
                        with zf.open(filename) as f:
                            with NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_f:
                                tmp_f.write(f.read())
                                tmp_p = tmp_f.name
                            
                            # Рекурсивно читаем файл из архива
                            if filename.lower().endswith('.docx'):
                                paragraphs.extend(self._read_docx(tmp_p, silent=True))
                            elif filename.lower().endswith('.pdf'):
                                paragraphs.extend(self._read_pdf(tmp_p))
                            elif filename.lower().endswith('.xlsx'):
                                paragraphs.extend(self._read_xlsx(tmp_p, silent=True))
                            elif filename.lower().endswith('.doc'):
                                paragraphs.extend(self._read_doc(tmp_p))
                            
                            os.unlink(tmp_p)
        except Exception:
            # Если это не zip или ошибка чтения, возвращаем пустоту
            pass
            
        return paragraphs

    def _get_223fz_notice_id(self, number):
        # метод имитирует поиск на сайте закупок для получения noticeInfoId (для 223-ФЗ)

        search_url = f"https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString={number}&morphology=on&pageNumber=1&sortDirection=false&recordsPerPage=_10&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz223=on&af=on&ca=on&pc=on&pa=on"
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}

        max_repeats = 3
        for attempt in range(max_repeats):
            try:
                logging.info(f"Попытка {attempt + 1} поиска noticeInfoId для 223-ФЗ (номер {number})...")
                response = requests.get(search_url, headers=headers, timeout=15)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, 'html.parser')

                results = soup.find_all('div', class_='search-registry-entry-block')
                if not results:
                    # попробуем другой класс, если сайт обновился
                    results = soup.find_all('div', class_='registry-entry__form')

                for res in results:
                    # ищем ссылку с классом m-0 или просто ссылку в заголовке
                    link_node = res.find('a', class_='m-0') or res.find('a', target='_blank')
                    if link_node and link_node.get('href'):
                        href = link_node['href']
                        # вытаскиваем noticeInfoId из href
                        match = re.search(r'noticeInfoId=(\d+)', href)
                        if match:
                            return match.group(1)

                logging.warning(f"Результаты поиска пусты для номера {number} (попытка {attempt + 1})")
            except Exception as e:
                logging.error(f"Ошибка при поиске noticeInfoId (попытка {attempt + 1}): {e}")

        return None




    def _check_keywords_in_text(self, text, keywords):
        #метод проверяет, содержит ли текст хотя бы одно из ключевых слов

        text_lower = text.lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                return True
        return False

    def _extract_date(self, text):
        # метод извлекает дату из текста (поддерживает форматы 31.12.2026 и "31" декабря 2026)
        if not text:
            return None

        # 1. Поиск стандартного формата DD.MM.YYYY
        date_pattern = r'\d{2}\.\d{2}\.\d{2,4}'
        dates = re.findall(date_pattern, text)
        if dates:
            return dates[0]

        # 2. Поиск текстового формата "31 декабря 2026" (с поддержкой кавычек и спецсимволов)
        months = [
            'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
            'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
        ]
        month_pattern = '|'.join(months)
        
        # Регулярка теперь учитывает возможные кавычки вокруг числа
        text_date_pattern = rf'["«»\'\s]*(\d{{1,2}})["«»\'\s]*\s+({month_pattern})\s+(\d{{4}})'
        
        match = re.search(text_date_pattern, text, re.IGNORECASE)
        if match:
            day, month_str, year = match.groups()
            # Превращаем в стандартный формат DD.MM.YYYY
            month_idx = months.index(month_str.lower()) + 1
            return f"{int(day):02d}.{month_idx:02d}.{year}"

        return None


    def _group_paragraphs(self, all_paragraphs):
        # Метод распределяет текст по группам, используя семантический анализ или ключевые слова
        
        # 1. Разбиваем все абзацы на предложения
        all_sentences = []
        for p in all_paragraphs:
            # Очищаем абзац от лишних пробелов и спецсимволов
            p = re.sub(r'\s+', ' ', p).strip()
            
            # Улучшенное разбиение
            sentences = re.split(r'(?<![гв])(?<!т\.д)(?<!и\.о)(?<=[.!?])\s+', p)
            
            for s in sentences:
                s = s.strip()
                # Фильтруем:
                # - слишком короткие (< 20)
                # - слишком длинные (> 500), скорее всего это склеенный мусор
                # - содержащие пустые поля для заполнения (____ или ....)
                if 20 < len(s) < 500 and not re.search(r'_{3,}|\.{4,}', s):
                    all_sentences.append(s)

        if not all_sentences:
            logging.warning("После очистки и разбиения не осталось подходящих предложений для анализа.")
            return

        # Удаляем дубликаты предложений (часто встречаются в разных документах)
        all_sentences = list(dict.fromkeys(all_sentences))
        logging.info(f"Для анализа подготовлено {len(all_sentences)} уникальных предложений.")

        # 2. Если есть семантический анализатор, используем его
        if self.semantic_analyzer:
            logging.info(f"Запуск семантического анализа...")
            for group_name in self.groups_config.keys():
                if group_name == "Срок действия контракта":
                    continue
                
                best_sentence = self.semantic_analyzer.find_best_sentence(all_sentences, group_name)
                
                if best_sentence:
                    self.groups[group_name] = [best_sentence]
                    # Специальная логика для даты
                    if group_name == "Примечание к сроку действия контракта":
                        date = self._extract_date(best_sentence)
                        if date:
                            self.contract_date = date
                else:
                    # Если семантика не помогла, пробуем ключевые слова (с более строгим фильтром)
                    self._fallback_keyword_search(all_sentences, group_name)
        else:
            # Только ключевые слова
            for group_name in self.groups_config.keys():
                if group_name == "Срок действия контракта":
                    continue
                self._fallback_keyword_search(all_sentences, group_name)

    def _fallback_keyword_search(self, sentences, group_name):
        """Запасной метод поиска по ключевым словам."""
        keywords = self.groups_config.get(group_name, [])
        best_match = None
        max_keywords = 0
        
        for s in sentences:
            # Считаем, сколько ключевых слов из конфига есть в этом предложении
            match_count = sum(1 for kw in keywords if kw.lower() in s.lower())
            
            if match_count > max_keywords:
                max_keywords = match_count
                best_match = s
        
        # Если нашли предложение хотя бы с одним ключевым словом
        if best_match and max_keywords > 0:
            logging.info(f"Фолбэк: для '{group_name}' выбрано по ключевым словам (совпало {max_keywords}): {best_match[:100]}...")
            self.groups[group_name] = [best_match]
            if group_name == "Примечание к сроку действия контракта":
                date = self._extract_date(best_match)
                if date:
                    self.contract_date = date

    def _process_contract_date(self):
        #метод обрабатывает специальную группу "Срок действия контракта"

        if self.contract_date:
            self.groups["Срок действия контракта"] = [self.contract_date]
        else:
            self.groups["Срок действия контракта"] = []



def get_args():
    #Настройка аргументов командной строки
    ap = argparse.ArgumentParser(description="Парсер для закупок")
    ap.add_argument('--mode', required=True, choices=['file', 'db'],
                    help='Режим ввода данных. Из файла или базы данных.')
    ap.add_argument('--input', type=str, help='Путь к файлу с номерами закупок (для режима file)')
    return ap.parse_args()


def read_numbers_from_file(file_path):
    """Чтение данных из файла. Поддерживает форматы:
    1. Просто номер закупки
    2. Номер закупки и ссылка через пробел/табуляцию
    """
    if not os.path.exists(file_path):
        logging.error(f"Файл {file_path} не найден.")
        return []
    
    results = []
    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            
            parts = line.split()
            number = parts[0]
            url = parts[1] if len(parts) > 1 else None
            results.append({'number': number, 'url': url})
            
    logging.info(f"Из файла {file_path} загружено {len(results)} записей.")
    return results


def get_db_connection():
    #Создает подключение к БД, используя параметры из .env
    try:
        conn_str = (
            "DRIVER={ODBC Driver 18 for SQL Server};"
            f"SERVER={os.getenv('DB_SERVER')};"
            f"DATABASE={os.getenv('DB_DATABASE')};"
            f"UID={os.getenv('DB_USERNAME')};"
            f"PWD={os.getenv('DB_PASSWORD')};"
            "TrustServerCertificate=yes;"
        )
        return pyodbc.connect(conn_str)
    except Exception as e:
        logging.error(f"Ошибка при подключении к БД: {e}")
        return None


def get_numbers_from_db():
    #Получение номеров и ссылок закупок из БД
    logging.info("Попытка получения данных из БД...")
    
    conn = get_db_connection()
    if not conn:
        logging.error("Подключение к БД не удалось, работа в режиме БД невозможна.")
        return []
    
    try:
        cursor = conn.cursor()

        # Фильтруем те закупки, где l.PlanTVal еще не заполнен (IS NULL)
        query = """SELECT DISTINCT t.NotifNr, t.SrcInf
                    FROM [Cursor].[dbo].[Tender] t
                    INNER JOIN [Cursor].[dbo].[Lot] l (nolock) on l.Tender_id = t.tender_id
                    INNER JOIN [Cursor].[dbo].[LotSpec] ls (nolock) on ls.lot_id = l.lot_id
                    WHERE ((FZ_ID = 44 AND NotifNr NOT LIKE '[a-zA-Z]%' AND len(NotifNr) = 19) 
                    OR (len(NotifNr) = 11 and NotifNr like '3%')) 
                    AND t.SYSDATE >= DATEADD(minute, -90, GETDATE()) 
                    AND l.PlanTVal IS NULL"""
        cursor.execute(query)
        results = [{'number': row[0], 'url': row[1]} for row in cursor.fetchall()]
        conn.close()
        logging.info(f"Из БД загружено {len(results)} записей для обработки.")
        return results
    except Exception as e:
        logging.error(f"Ошибка при выполнении запроса к БД: {e}")
        return []


def save_to_db(number, data):
    #Сохранение результатов в БД
    logging.info(f"Сохранение результатов для закупки {number} в БД...")
    
    conn = get_db_connection()
    if not conn:
        logging.error(f"Подключение к БД не удалось. Данные для {number} не сохранены.")
        return
    
    try:
        cursor = conn.cursor()
        
        # Подготавливаем текстовые данные из найденных групп (объединяем абзацы)
        def get_field_val(group_name):
            val = "\n".join(data.get(group_name, []))
            return val if val.strip() else "Нет данных"

        fs = get_field_val("Источник финансирования")
        sl = get_field_val("Требования к сроку годности")
        pt = get_field_val("Порядок оплаты товаров")
        dt = get_field_val("Условия поставки")
        cdn = get_field_val("Примечание к сроку действия контракта")
        cd = get_field_val("Срок действия контракта")

        # 1. Обновляем таблицу [Tender]
        query_tender = """
        UPDATE [Cursor].[dbo].[Tender]
        SET TenderDocReglament = ?, RequirementToExpiryDate = ?
        WHERE NotifNr = ?
        """
        cursor.execute(query_tender, (fs, sl, number))
        
        # 2. Обновляем таблицу [Lot]
        # Обновляем все лоты, связанные с этим номером закупки (через Tender_id)
        query_lot = """
        UPDATE l
        SET l.PaymentReglament = ?, 
            l.PlanTVal = ?, 
            l.ContrExpVal = ?, 
            l.SupplyDt = ?
        FROM [Cursor].[dbo].[Lot] l
        INNER JOIN [Cursor].[dbo].[Tender] t ON l.Tender_id = t.tender_id
        WHERE t.NotifNr = ?
        """
        cursor.execute(query_lot, (pt, dt, cdn, cd, number))
        
        conn.commit()
        conn.close()
        logging.info(f"Данные для {number} успешно сохранены в таблицы Tender и Lot.")
    except Exception as e:
        logging.error(f"Ошибка при сохранении в БД для {number}: {e}")


def main():
    logging.info("Запуск парсера.")
    args = get_args()
    
    # Инициализируем семантический анализатор один раз при запуске
    analyzer = None
    excel_path = "Данные из базы.xlsx"
    cache_path = "embeddings_cache.pt"
    
    if os.path.exists(excel_path) or os.path.exists(cache_path):
        analyzer = SemanticAnalyzer(excel_path, cache_path=cache_path)
    else:
        logging.warning(f"Ни файл {excel_path}, ни кеш {cache_path} не найдены. Семантический поиск отключен.")

    parser = DocumentParser(semantic_analyzer=analyzer)
    
    numbers = []
    
    if args.mode == 'file':
        if not args.input:
            logging.error("Для режима 'file' необходимо указать путь к файлу через --input")
            return
        numbers = read_numbers_from_file(args.input)
    elif args.mode == 'db':
        numbers = get_numbers_from_db()

    if not numbers:
        logging.warning("Список номеров для обработки пуст. Завершение работы.")
        return

    logging.info(f"Всего предстоит обработать закупок: {len(numbers)}")

    for entry in numbers:
        number = entry['number']
        url = entry['url']
        try:
            # Сбрасываем группы перед каждым новым парсингом
            parser.groups = {group: [] for group in parser.groups_config.keys()}
            parser.contract_date = None
            
            # Передаем и номер, и ссылку в метод парсинга
            groups = parser.parse_by_number(number, provided_url=url)
            
            # Временный принт результатов для отладки
            print(f"\n--- РЕЗУЛЬТАТЫ ДЛЯ {number} ---")
            for group_name, sentences in groups.items():
                content = sentences[0] if sentences else "Нет данных"
                print(f"[{group_name}]: {content}")
            print("-" * 30 + "\n")
            
            # Сохранение в БД (временная заглушка)
            # save_to_db(number, groups)
            
            logging.info(f"Закупка {number} успешно обработана.")

        except Exception as e:
            logging.error(f"Критическая ошибка при обработке закупки {number}: {e}")

    logging.info("Работа парсера завершена.")

if __name__ == "__main__":
    main()
